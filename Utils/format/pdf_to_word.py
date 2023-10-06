from PyPDF2 import PdfReader
from docx import Document

def pdf_to_word(pdf_path, word_path):
    pdf_file = open(pdf_path, 'rb')
    pdf_reader = PdfReader(pdf_file)

    document = Document()

    for page in pdf_reader.pages:
        text = page.extract_text()
        paragraphs = text.split('\n')

        for paragraph in paragraphs:
            if paragraph:
                document.add_paragraph(paragraph)
        document.add_paragraph('\n')
    document.save(word_path)
    pdf_file.close()

# 使用示例
pdf_path = '.pdf'
word_path = 'res.docx'

pdf_to_word(pdf_path, word_path)